import tkinter as tk
from tkinter import ttk
from docx import Document
window = tk.Tk()
window.title("VTU CALCULATOR")
window.geometry("700x500")
window['background']='#BFEAF5'
def calculator():
    def input_cgpa():
        global name, usn, branch, college, phone
        name = name_entry.get()
        usn = usn_entry.get()
        branch = Branch_combobox.get()
        college = college_entry.get()
        phone = Phone_entry.get()
        cgpa_interface()
    def cgpa_interface():
        global sem
        sem = int(sem_combobox.get())
        frame1.destroy()
        global s
        s=[]
        frame = tk.Frame(window)
        frame.pack(padx=20, pady=20)
        frame['background']='#BFEAF5'
        Marks_Details = tk.LabelFrame(frame, text="SGPA Inputs")
        Marks_Details.grid(row=0,column=0,padx=10, pady=10)
        Marks_Details['background']='#BFEAF5'
        for i in range(sem):
            sem_label = tk.Label(Marks_Details, text=f"Semester - {i+1}")
            sem_label.grid(row = i, column = 0, padx=10, pady=10)
            sem_label['background']='#F5F0BB'
            sgpa_entry = tk.Entry(Marks_Details)
            sgpa_entry.grid(row = i,column = 1, columnspan=2, padx=20, pady=10)
            s.append(sgpa_entry)
            cal_cgpa = tk.Button(frame, text="Calculate CGPA", command=display_cgpa,bg='#F5F0BB',activebackground='#abb072')
            cal_cgpa.grid(row=sem+1, column=0, columnspan=2, sticky="news")
    def display_cgpa():
            global val,cgpa
            val=[]
            for sgpa_entry in s:
                value = (sgpa_entry.get())
                val.append(float(value))
            sum_val = sum(val)
            op = sum_val/(len(s))
            cgpa = round(op,2)
            frame2 = tk.Frame(window)
            frame2.pack()
            frame2['background']='#BFEAF5'
            result_label1 = tk.Label(frame2, text=f"CGPA : \t{cgpa}")
            result_label1.grid(row=0,column=0,padx = 20, pady = 20)
            result_label1['background']='#F5F0BB'
            print_button=tk.Button(frame2,text="Print",command=print_docx_cgpa, bg='#F5F0BB',activebackground='#abb072')
            print_button.grid(row=1,column=0,sticky="news", padx=20, pady=10)
    def print_docx_cgpa():
            document = Document()
            document.add_heading('VISVESVARAYA TECHNOLOGICAL UNIVERSITY', 0)
            document.add_heading('VTU RESULTS OF UG EXAMINATION', level=1)
            #name, usn, branch, sem, college, phone
            details1= document.add_table(rows=2, cols=6)
            details1.style='Medium Shading 1 Accent 3'
            first_row1 = details1.rows[0].cells
            first_row1[0].text = 'NAME'
            first_row1[1].text = 'USN'
            first_row1[2].text = 'BRANCH'
            first_row1[3].text = 'SEMESTER'
            first_row1[4].text = 'COLLEGE'
            first_row1[5].text = 'PHONE'
            second_row1 = details1.rows[1].cells
            second_row1[0].text = str(name)
            second_row1[1].text = str(usn)
            second_row1[2].text = str(branch)
            second_row1[3].text = str(sem)
            second_row1[4].text = str(college)
            second_row1[5].text = str(phone)
            p1=document.add_paragraph('')
            table1=document.add_table(rows=1, cols=3)
            table1.style='Medium Shading 1 Accent 3'
            hd_cell= table1.rows[0].cells
            hd_cell[0].text = 'Sl no.'
            hd_cell[1].text = 'SEMESTER'
            hd_cell[2].text = 'SGPA'
            for i in range(sem):
                row_cell = table1.add_row().cells
                row_cell[0].text = str(i+1)
                row_cell[1].text = "Semester - "+str(i+1)
                row_cell[2].text = str(val[i])
            p1=document.add_paragraph('')
            table2 = document.add_table(rows=1, cols=2)
            table2.style = 'Medium Shading 1'
            cell = table2.rows[0].cells
            cell[0].text = 'CGPA'
            cell[1].text = str(cgpa)
            document.save(f'student-{usn}-result.docx')
    def input_sgpa():
        global name, usn, branch, college, phone
        name = name_entry.get()
        usn = usn_entry.get()
        branch = Branch_combobox.get()
        college = college_entry.get()
        phone = Phone_entry.get()
        sgpa_interface()
    def sgpa_interface():
        global data_entries, externals, internals, SGPA, entry_num_inputs, frame3
        global sem
        sem = int(sem_combobox.get())
        frame1.destroy()
        frame3 = tk.Frame(window)
        frame3.pack(padx=10, pady=10)
        frame3['background']='#BFEAF5'
        data_entries = []
        externals=[]
        internals=[]
         # Create a label and entry for specifying the number of data inputs
        label_num_inputs = tk.Label(frame3, text="Enter number of Subjects:")
        label_num_inputs.grid(row=0,column=0, padx=20, pady=20)
        label_num_inputs['background']='#F5F0BB'
        entry_num_inputs = tk.Entry(frame3)
        entry_num_inputs.grid(row=0, column=1, padx=20, pady=10)
        # Create a button to create the data inputs
        create_button = tk.Button(frame3, text="Create Inputs", command=create_data_inputs,bg='#F5F0BB',activebackground='#abb072')
        create_button.grid(row=1,column=1, padx=20, pady=10)
    def create_data_inputs():
        global num_inputs, subjects
        num_inputs = int(entry_num_inputs.get())
        subjects=[]
        for i in range(num_inputs):
            label1 = tk.Label(frame3, text="Enter Subject code:")
            label1.grid(row=3,column=0, pady=10)
            label1['background']='#F5F0BB'
            

            label2= tk.Label(frame3, text="Enter Credits :")
            label2.grid(row=3,column=1, pady=10)
            label2['background']='#F5F0BB'
            

            label3 = tk.Label(frame3, text="Enter Internal Marks:")
            label3.grid(row=3,column=2, padx=5, pady=10)
            label3['background']='#F5F0BB'

            label4 = tk.Label(frame3, text="Enter External Marks:")
            label4.grid(row=3,column=3, padx=5, pady=10)
            label4['background']='#F5F0BB'
            
            subject_entry = tk.Entry(frame3)
            subject_entry.grid(row=i+4,column=0, padx=5, pady=10)
            subjects.append(subject_entry)
            
            credits = tk.Entry(frame3)
            credits.grid(row=i+4,column=1, padx=5, pady=10)
            data_entries.append(credits)
            
            external= tk.Entry(frame3)
            external.grid(row=i+4,column=2, padx=5, pady=10)
            externals.append(external)

            internal = tk.Entry(frame3)
            internal.grid(row=i+4,column=3, padx=5, pady=10)
            internals.append(internal)

            result_button = tk.Button(frame3, text="Result", command=save_data,bg='#F5F0BB',activebackground='#abb072')
            result_button.grid(row=num_inputs+5,column=2,columnspan=2, padx=10, pady=10,sticky="news")  
    def save_data():
        global sub, data, ext, inte, marks,SGPA
        data = []
        ext=[]
        inte=[]
        sub=[]
        marks=[]
        form=0
        sub_credits=[]
        total_credits=0
        for subject_entry in subjects:
            value4= subject_entry.get()
            sub.append(value4)
        print(sub)  
        for credits in data_entries:
            value1 = credits.get()
            data.append(int(value1))
            # Perform any required processing or storage operations with the data
        s=sum(data)
        
        for external in externals:
            value2 = external.get()
            ext.append(int(value2))
         # Perform any required processing or storage operations with the data
        x=sum(ext)
        

        for internal in internals:
            value3 = internal.get()
            inte.append(int(value3))
         # Perform any required processing or storage operations with the data
        y=sum(inte)
        marks=[sum(i) for i in zip(ext,inte)]
        print(marks)
        total_marks_obtained = sum(marks)
        a=len(ext)
        percentage=(total_marks_obtained/(a*100))*100
        for i in range (0,len(marks)):
            if ext[i]<=50 and inte[i]<=50:
                if marks[i]==100:
                    form=(10*data[i])
                elif marks[i]<100:
                    form=((marks[i]//10)+1)*data[i]
                sub_credits.append(form)
                total_credits=int(total_credits+form)
        SGPA=round((total_credits/sum(data)),4)
        display(SGPA,total_marks_obtained,percentage)       
        print(f"SGPA={SGPA}")
    def display(SGPA,total_marks_obtained,percentage):
        frame3.destroy()
        frame4 = tk.Frame(window)
        frame4.pack(padx=20,pady=50)
        frame4['background']='#BFEAF5'
        result_label1 = tk.Label(frame4, text=f"SGPA : \t{SGPA}")
        result_label1.grid(row=0,column=0,padx = 20, pady = 10)
        result_label1['background']='#F5F0BB'
        result_label2 = tk.Label(frame4, text=f"TOTAL MARKS : \t{total_marks_obtained}")
        result_label2.grid(row=1,column=0,padx = 20, pady = 10)
        result_label2['background']='#F5F0BB'
        result_label3 = tk.Label(frame4, text=f"PERCENTAGE : \t{percentage}")
        result_label3.grid(row=2,column=0,padx = 20, pady = 10)
        result_label3['background']='#F5F0BB'
        print_button=tk.Button(frame4,text="Print",command=print_docx_sgpa, bg='#F5F0BB',activebackground='#abb072')
        print_button.grid(row=3,column=0,sticky="news", padx=20, pady=10)
    def print_docx_sgpa():
        document = Document()
        document.add_heading('VISVESVARAYA TECHNOLOGICAL UNIVERSITY', 0)
        document.add_heading('VTU RESULTS OF UG EXAMINATION', level=1)
        #name, usn, branch, sem, college, phone
        details1= document.add_table(rows=2, cols=6)
        details1.style='Medium Shading 1 Accent 3'
        first_row1 = details1.rows[0].cells
        first_row1[0].text = 'NAME'
        first_row1[1].text = 'USN'
        first_row1[2].text = 'BRANCH'
        first_row1[3].text = 'SEMESTER'
        first_row1[4].text = 'COLLEGE'
        first_row1[5].text = 'PHONE'
        second_row1 = details1.rows[1].cells
        second_row1[0].text = name
        second_row1[1].text = usn
        second_row1[2].text = branch
        second_row1[3].text = str(sem)
        second_row1[4].text = college
        second_row1[5].text = phone
        p1=document.add_paragraph('')

        #marks table
        table = document.add_table(rows=1, cols=5)
        table.style='Medium Shading 1 Accent 3'
        hd_cells = table.rows[0].cells
        hd_cells[0].text = 'SUBJECT CODE'
        hd_cells[1].text = 'INTERNAL MARKS'
        hd_cells[2].text = 'EXTERNAL MARKS'
        hd_cells[3].text = 'TOTAL MARKS'
        hd_cells[4].text = 'RESULT'

        for i in range(num_inputs):
            row_cells = table.add_row().cells
            row_cells[0].text = str(sub[i])
            row_cells[1].text = str(inte[i])
            row_cells[2].text = str(ext[i])
            row_cells[3].text = str(marks[i])
            if(inte[i]>=20 and ext[i]>=18):
                result = 'P'
            else:
                result = 'F'
            row_cells[4].text = result
        p1=document.add_paragraph('')
        result_sgpa = document.add_table(rows=2,cols=3)
        result_sgpa.style='Medium Shading 1'
        results_row1= result_sgpa.rows[0].cells
        results_row1[0].text = 'SGPA'
        results_row1[1].text = 'TOTAL MARKS'
        results_row1[2].text = 'PERCENTAGE'
        results_row2= result_sgpa.rows[1].cells
        results_row2[0].text = str(SGPA)
        results_row2[1].text = str(sum(marks))
        results_row2[2].text = f"{sum(marks)/num_inputs} %"
        document.save(f'student-{usn}-result.docx')
    #Student Details
    frame1 = tk.Frame(window)
    frame1.pack(pady=50)
    frame1['background']='#BFEAF5'
    Student_Details = tk.LabelFrame(frame1, text="Student Details")
    Student_Details.grid(row=0, column=0, padx=10, pady=10)
    Student_Details['background']='#BFEAF5'
    Name_label = tk.Label(Student_Details, text="Name")
    Name_label.grid(row=0, column=0, padx=20, pady=10)
    Name_label['background']='#F5F0BB'
    USN_Label = tk.Label(Student_Details, text="USN")
    USN_Label.grid(row=1, column=0, padx=20, pady=10)
    USN_Label['background']='#F5F0BB'
    Branch_Label = tk.Label(Student_Details, text="Branch")
    Branch_Label.grid(row=2, column=0, padx=20, pady=10)
    Branch_Label['background']='#F5F0BB'
    Sem_Label = tk.Label(Student_Details, text="Semester")
    Sem_Label.grid(row=3, column=0, padx=20, pady=10)
    Sem_Label['background']='#F5F0BB'
    College_Label = tk.Label(Student_Details, text="College")
    College_Label.grid(row=4, column=0, padx=20, pady=10)
    College_Label['background']='#F5F0BB'
    Phone_Label = tk.Label(Student_Details, text="Phone")
    Phone_Label.grid(row=5, column=0, padx=20, pady=10)
    Phone_Label['background']='#F5F0BB'
    #entries
    name_entry = tk.Entry(Student_Details)
    name_entry.grid(row=0, column=1, columnspan=2, padx=20, pady=10)
    usn_entry = tk.Entry(Student_Details)
    usn_entry.grid(row=1, column=1, columnspan=2, padx=10, pady=10)
    Branch_combobox = ttk.Combobox(Student_Details,
    values=[
        "ARTIFICIAL INTELLIGENCE & MACHINE LEARNING", "CIVIL ENGINEERING",
        "COMPUTER SCIENCE & ENGINEERING", "INFORMATION SCIENCE & ENGINEERING","ELECTRICAL AND ELECTRONICS ENGINEERING",
        "ELECTRONICS & COMMUNICATION ENGINEERING", "MECHANICAL ENGINEERING","INDUSTRIAL ENGINEERING AND MANAGEMENT",
        "ELECTRONICS AND INSTRUMENTATION ENGINEERING","ELECTRONICS AND TELECOMMUNICATION ENGINEERING"
    ])
    Branch_combobox.grid(row=2, column=1, padx=20, pady=20)
    sem_combobox = ttk.Combobox(Student_Details, values=[1, 2, 3, 4, 5, 6, 7, 8])
    sem_combobox.grid(row=3, column=1, padx=10, pady=10)
    college_entry = tk.Entry(Student_Details)
    college_entry.grid(row=4, column=1, columnspan=2, padx=10, pady=10)
    Phone_entry = tk.Entry(Student_Details)
    Phone_entry.grid(row=5, column=1, columnspan=2, padx=10, pady=10) 
    calc_cgpa=tk.Button(frame1,text="Calculate CGPA",command=input_cgpa,bg='#F5F0BB',activebackground='#abb072')
    calc_cgpa.grid(row=6,column=0,sticky="news")
    calc_sgpa=tk.Button(frame1,text="Calculate SGPA",bg='#F5F0BB', command=input_sgpa,activebackground='#abb072')
    calc_sgpa.grid(row=7,column=0,sticky="news")


calculator()
window.mainloop()
